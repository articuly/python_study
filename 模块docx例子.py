# coding:utf-8
from docx import Document

d = Document()
d.add_paragraph('Life is short, You need python.')
d.add_picture('articuly.jpg')
d.save('python sample.docx')
